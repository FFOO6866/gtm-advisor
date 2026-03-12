#!/usr/bin/env python3
"""Seed demo data for Hi Meet.AI by running the production agent pipeline.

Creates:
- User: demo@himeetai.com / himeetai2024
- Company: Hi Meet.AI
- Runs the full 6-agent GTM analysis pipeline (Market Intel, Competitor,
  Customer Profiler, Lead Hunter, Campaign Architect) with real data sources.
- Leads, market insights, and campaign are persisted by the agents themselves.

Usage:
    uv run python scripts/seed_himeetai.py
"""

import asyncio
import sys
from pathlib import Path
from uuid import uuid4

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

BIZ_PLAN_PATH = Path.home() / "Downloads" / "20260207 Biz Plan_05.docx"


def _read_biz_plan() -> str | None:
    """Extract raw text from the business plan docx."""
    if not BIZ_PLAN_PATH.exists():
        print(f"  ⚠ Business plan not found at {BIZ_PLAN_PATH} — proceeding without it.")
        return None
    try:
        from docx import Document
        doc = Document(str(BIZ_PLAN_PATH))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts)
        print(f"  → Business plan loaded: {len(text):,} characters from {BIZ_PLAN_PATH.name}")
        return text
    except Exception as e:
        print(f"  ⚠ Could not read business plan: {e}")
        return None

from sqlalchemy import select

from packages.core.src.types import IndustryVertical
from packages.database.src.models import (
    Analysis,
    AnalysisStatus,
    Company,
    SubscriptionTier,
    User,
)
from packages.database.src.session import async_session_factory, close_db, init_db
from services.gateway.src.auth.utils import get_password_hash
from services.gateway.src.routers.analysis import AnalysisRequest, run_analysis


def _build_request(document_text: str | None) -> AnalysisRequest:
    return AnalysisRequest(
        company_name="Hi Meet.AI",
        website="https://himeetai.com",
        description=(
            "Hi Meet.AI is an AI-native growth firm that replaces the traditional "
            "manpower-heavy agency model with a proprietary multi-agent architecture. "
            "It provides on-demand CMO, Strategy Director, Visual Architect, and Data "
            "Scientist agentic swarms that manage the full growth lifecycle — from "
            "competitive intelligence and GTM strategy to content, outreach, and "
            "performance optimisation — at machine speed."
        ),
        industry=IndustryVertical.PROFESSIONAL_SERVICES,
        goals=[
            "Establish Hi Meet.AI as Singapore's leading AI-native growth firm",
            "Win retainer clients across three tiers: Foundation SMEs, Core Growth mid-market, Enterprise",
            "Demonstrate measurable ROI superiority over legacy agencies",
            "Expand to APAC markets after proving Singapore model",
        ],
        challenges=[
            "Educating market on agentic AI vs traditional agency model",
            "Building trust and social proof as a new category entrant",
            "Competing against established agency brands with existing client relationships",
            "Pricing the proprietary multi-agent architecture vs hourly agency billing",
        ],
        competitors=[
            "Traditional marketing agencies (Publicis, Dentsu, Ogilvy Singapore)",
            "Digital growth agencies (FirstPage, MediaOne)",
            "AI marketing tools (Jasper, Copy.ai, HubSpot AI)",
        ],
        target_markets=[
            "Singapore SMEs and startups seeking enterprise-level GTM resources",
            "Mid-market firms scaling APAC operations",
            "Enterprise brands looking to replace legacy agency model",
        ],
        value_proposition=(
            "Hi Meet.AI delivers the strategic bench strength of a full agency team "
            "through AI agents that execute in minutes rather than months, at a fraction "
            "of traditional agency cost — with measurable, attribution-tracked outcomes."
        ),
        additional_context=document_text,
        include_market_research=True,
        include_competitor_analysis=True,
        include_customer_profiling=True,
        include_lead_generation=True,
        include_campaign_planning=True,
        lead_count=25,
    )


async def main() -> None:
    print("Initialising database…")
    await init_db()

    # ── Load business plan document ───────────────────────────────────────────
    print("\nLoading business plan document…")
    document_text = _read_biz_plan()
    HIMEETAI_REQUEST = _build_request(document_text)

    async with async_session_factory() as db:
        # ── 1. User ──────────────────────────────────────────────────────────
        existing = await db.execute(select(User).where(User.email == "demo@himeetai.com"))
        user = existing.scalar_one_or_none()

        if user:
            print(f"User already exists: {user.email}")
            company_q = await db.execute(
                select(Company).where(Company.owner_id == user.id)
            )
            company = company_q.scalar_one_or_none()
            if not company:
                print("No owned company found — will create one.")
                company = None
            else:
                print(f"Found company: {company.name} ({company.id})")
                # Update context_sources on existing company with biz plan
                if document_text:
                    from services.gateway.src.routers.analysis import _make_source_entry
                    sources = list(company.context_sources or [])
                    sources = [s for s in sources if not (s.get("type") == "document" and s.get("name") == "uploaded_document")]
                    sources.append(_make_source_entry("document", "uploaded_document", document_text))
                    company.context_sources = sources
                    print("  → Updated context_sources on existing company")
        else:
            print("Creating user demo@himeetai.com …")
            user = User(
                email="demo@himeetai.com",
                hashed_password=get_password_hash("himeetai2024"),
                full_name="Hi Meet.AI Demo",
                company_name="Hi Meet.AI",
                tier=SubscriptionTier.TIER1,
                is_active=True,
            )
            db.add(user)
            await db.flush()
            company = None

        # ── 2. Company ───────────────────────────────────────────────────────
        if not company:
            print("Creating company Hi Meet.AI …")
            company = Company(
                owner_id=user.id,
                name="Hi Meet.AI",
                website="https://himeetai.com",
                description=HIMEETAI_REQUEST.description,
                industry="professional_services",
                headquarters="Singapore",
                employee_count="11-50",
                funding_stage="Seed",
                goals=HIMEETAI_REQUEST.goals,
                challenges=HIMEETAI_REQUEST.challenges,
                competitors=HIMEETAI_REQUEST.competitors,
                target_markets=HIMEETAI_REQUEST.target_markets,
                value_proposition=HIMEETAI_REQUEST.value_proposition,
                products=[
                    {
                        "name": "Offering A — Strategic Foundation",
                        "description": "Entry retainer: AI-powered competitive intelligence, ICP refinement, GTM advisory, messaging strategy",
                        "type": "retainer",
                    },
                    {
                        "name": "Offering B — Integrated Growth Partner",
                        "description": "Core retainer: quarterly GTM strategy, campaign orchestration, strategic comms, AI optimisation, ROI reporting",
                        "type": "retainer",
                    },
                    {
                        "name": "Offering C — Embedded Strategic Partner",
                        "description": "Enterprise retainer: boardroom advisory, deep market intelligence, multi-vendor orchestration, pipeline acceleration",
                        "type": "retainer",
                    },
                ],
                context_sources=(
                    [{"type": "document", "name": "uploaded_document", "text": document_text, "chars": len(document_text), "added_at": "2026-03-11T00:00:00+00:00"}]
                    if document_text else []
                ),
                enrichment_confidence=0.0,
            )
            db.add(company)
            await db.flush()
            print(f"  → Company ID: {company.id}")

        # ── 3. Analysis record ────────────────────────────────────────────────
        analysis_id = uuid4()
        analysis = Analysis(
            id=analysis_id,
            company_id=company.id,
            user_id=user.id,
            status=AnalysisStatus.PENDING,
            progress=0,
            completed_agents=[],
        )
        db.add(analysis)
        # Attach company_id to request so agents write to the right company
        HIMEETAI_REQUEST.company_id = company.id
        await db.commit()

    # ── 4. Run the production agent pipeline ─────────────────────────────────
    print(f"\nRunning production agents for analysis {analysis_id} …")
    print("This will take 3–8 minutes (Perplexity + EODHD + LLM calls).\n")

    await run_analysis(analysis_id, HIMEETAI_REQUEST)

    # ── 5. Report results ─────────────────────────────────────────────────────
    async with async_session_factory() as db:
        analysis = await db.get(Analysis, analysis_id)
        leads_count = len(analysis.leads or [])
        insights_count = len(analysis.market_insights or [])
        has_campaign = bool(analysis.campaign_brief)
        competitors_count = len(analysis.competitor_analysis or [])
        personas_count = len(analysis.customer_personas or [])

    await close_db()

    print("\nDone!")
    print("─" * 50)
    print(f"Status:      {analysis.status.value}")
    print(f"Confidence:  {(analysis.total_confidence or 0):.0%}")
    print(f"Agents:      {', '.join(analysis.agents_used or [])}")
    print(f"Leads:       {leads_count}")
    print(f"Insights:    {insights_count}")
    print(f"Competitors: {competitors_count}")
    print(f"Personas:    {personas_count}")
    print(f"Campaign:    {'yes' if has_campaign else 'no'}")
    print("─" * 50)
    print("Login:    demo@himeetai.com")
    print("Password: himeetai2024")
    print("─" * 50)


if __name__ == "__main__":
    asyncio.run(main())
