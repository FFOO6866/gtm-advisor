"""Document parsing — extract company profile fields from uploaded files.

Accepts PDF or plain-text files (max 10 MB). Extracts raw text, then uses
GPT-4o-mini with temperature=0 and a strict "extract only, do not invent"
system prompt to return structured company-profile fields.

Edge cases handled:
  - Image-only / scanned PDFs   → 422 with clear message
  - Spoofed file type           → validated by actually parsing, not MIME
  - Very long documents         → text truncated to MAX_TEXT_CHARS before LLM
  - LLM extraction failure      → 502 with advice to fill manually
  - Empty / corrupt files       → 400 / 422 with explanation
"""

from __future__ import annotations

import io
import json
from typing import Any

import structlog
from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel, Field

logger = structlog.get_logger()

router = APIRouter(prefix="/api/v1/documents", tags=["documents"])

MAX_FILE_BYTES = 10 * 1024 * 1024   # 10 MB
MAX_TEXT_CHARS = 100_000             # ~25 k tokens — GPT-4o handles up to 128 k

VALID_INDUSTRIES = {
    "fintech", "saas", "ecommerce", "healthcare",
    "logistics", "professional_services", "education", "other",
}

SYSTEM_PROMPT = """\
You are an independent business analyst preparing a company profile brief for a GTM strategy
engagement. The uploaded document was written BY the company (website copy, pitch deck, or
business plan). It may use first-person ("we deliver…") or second-person marketing language
("transform your business…"). IMPORTANT: in marketing copy "your" and "you" refer to the
document company's CLIENTS, not to the company itself.

Your task: extract factual information ABOUT the company and write it in clear, neutral,
third-person analytical language. Never copy marketing taglines verbatim.

Return ONLY a valid JSON object with exactly these keys:

{
  "company_name": <string or null — the company's trading name>,

  "description": <string or null — 2-4 sentences covering: (1) what the company does,
    (2) who its customers are, (3) how it delivers its service or product, as stated in
    the document. Use neutral, factual language — no superlatives. Do not include
    outcomes the company promises to clients; describe the company's own offering.
    Example for a professional services firm: "Meridian Advisory is a growth consultancy
    that helps mid-market technology companies enter new geographic markets. It provides
    embedded teams of market strategists and sales specialists who work alongside client
    leadership to design and execute market entry plans.">,

  "industry": <one of: "fintech" | "saas" | "ecommerce" | "healthcare" | "logistics"
    | "professional_services" | "education" | "other" — or null if unclear>,

  "value_proposition": <string or null — one sentence describing the core benefit the
    company delivers to its customers, as stated or demonstrated in the document.
    Write as an analyst observation: state what the company does and the outcome for
    clients. Do NOT use "you/your/we/our".
    Example: "Meridian Advisory enables technology companies to compress their market
    entry timeline by providing on-the-ground commercial expertise in target markets.">,

  "goals": <string array, max 5 — the company's own stated business or growth objectives,
    exactly as the document describes them for the company itself (not promises to clients).
    Include both quantitative targets and directional objectives if present.
    Empty array if no goals are explicitly stated for the company.>,

  "challenges": <string array, max 5 — obstacles the company explicitly acknowledges
    it faces: competitive pressure, market conditions, operational constraints, etc.
    Only include challenges the document directly states. Empty array if none stated.>,

  "competitors": <string array, max 5 — names of competing companies or products
    explicitly mentioned in the document. Empty array if none are named.>,

  "target_markets": <string array — geographic markets (countries or regions) the company
    serves or plans to enter, as stated in the document. Empty array if not specified.>
}

CRITICAL RULES:
- Never use first-person ("we", "our") or second-person ("you", "your") in any text output.
- Only include facts explicitly stated in the document — do not infer or invent anything.
- Describe what the company offers; do not list what it promises clients will achieve.
- If a field cannot be determined from the document, return null (strings) or [] (arrays).
- Return only the JSON object — no preamble, no markdown, no explanation.
"""

WEB_SYSTEM_PROMPT = """\
You are an independent business analyst preparing a company profile brief from a scraped company
website. The text below was extracted from the company's public website — it contains marketing
copy, navigation elements, testimonials, pricing pages, CTAs, and other web-specific content.

Your task: cut through the marketing language and extract FACTUAL information about the company.
Ignore testimonials, client logos, pricing CTAs, cookie notices, and navigation text.

Return ONLY a valid JSON object with exactly these keys:

{
  "company_name": <string or null — the company's trading name, not a tagline>,

  "description": <string or null — 2-4 sentences covering: (1) what the company does,
    (2) who its customers are, (3) how it delivers its service or product.
    Use neutral, factual language — strip marketing superlatives ("revolutionary",
    "world-class", "transform your business"). Do not include client testimonials or
    outcome claims. Focus on the actual product/service.
    Example: "Acme Corp provides cloud-based inventory management software for
    mid-market retailers in Southeast Asia. The platform integrates with POS systems
    and e-commerce platforms to provide real-time stock visibility.">,

  "industry": <one of: "fintech" | "saas" | "ecommerce" | "healthcare" | "logistics"
    | "professional_services" | "education" | "other" — or null if unclear>,

  "value_proposition": <string or null — one sentence describing the core benefit.
    Write as an analyst observation, not marketing copy.
    Example: "Acme Corp reduces stockout rates for retailers by providing real-time
    inventory synchronisation across physical and online channels.">,

  "goals": <string array, max 5 — business objectives if stated (e.g. "expanding to
    Indonesia", "launching enterprise tier"). Look in About/Mission/Vision sections.
    Empty array if none explicitly stated.>,

  "challenges": <string array, max 5 — obstacles the company faces, if mentioned.
    Often found in blog posts or press releases on the site. Empty array if none.>,

  "competitors": <string array, max 5 — competing companies mentioned on the site
    (comparison pages, "alternatives to" content). Empty array if none named.>,

  "target_markets": <string array — geographic markets served, from contact addresses,
    office locations, or explicit "we serve" statements. Empty array if not specified.>
}

CRITICAL RULES:
- Never use first-person ("we", "our") or second-person ("you", "your") in any output.
- Only include facts derivable from the page content — do not infer or invent.
- Ignore: testimonials, case study metrics, pricing details, cookie/privacy notices.
- If a field cannot be determined, return null (strings) or [] (arrays).
- Return only the JSON object — no preamble, no markdown, no explanation.
"""


# ─── Response model ────────────────────────────────────────────────────────

class ParsedCompanyProfile(BaseModel):
    company_name: str | None = None
    description: str | None = None
    industry: str | None = None
    value_proposition: str | None = None
    goals: list[str] = []
    challenges: list[str] = []
    competitors: list[str] = []
    target_markets: list[str] = []
    extracted_chars: int = 0
    extracted_text: str | None = None  # Truncated text sent to LLM, for agent context
    warning: str | None = None


# ─── Text extraction helpers ───────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        return "\n".join(pages)
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        raise ValueError(f"Could not read DOCX: {exc}") from exc


def _extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file — unsupported character encoding.")


def _extract_text(data: bytes, filename: str, content_type: str) -> str:
    fn = filename.lower()
    ct = content_type.lower()

    if fn.endswith(".pdf") or "pdf" in ct:
        return _extract_pdf(data)

    if fn.endswith(".docx") or "wordprocessingml" in ct or "msword" in ct:
        return _extract_docx(data)

    if fn.endswith(".txt") or "text/plain" in ct:
        return _extract_txt(data)

    # Unknown extension — try PDF, then DOCX, then plain text
    for extractor in (_extract_pdf, _extract_docx, _extract_txt):
        try:
            return extractor(data)
        except ValueError:
            continue
    raise ValueError("Could not extract text from this file. Please use PDF, DOCX, or TXT.")


# ─── Output coercion ───────────────────────────────────────────────────────

def _to_str_list(value: Any) -> list[str]:
    """Coerce LLM output to list[str], guarding against wrong types."""
    if isinstance(value, list):
        return [str(item) for item in value if item]
    if isinstance(value, str) and value.strip():
        # Model returned a delimited string instead of an array.
        # Prefer newline splitting (common for numbered lists); fall back to comma.
        sep = "\n" if "\n" in value else ","
        return [item.strip() for item in value.split(sep) if item.strip()]
    return []


# ─── LLM extraction ────────────────────────────────────────────────────────

async def _llm_extract(text: str, system_prompt: str = SYSTEM_PROMPT) -> dict[str, Any]:
    """Call GPT-4o-mini with structured output to extract company fields."""
    from openai import AsyncOpenAI

    client = AsyncOpenAI()
    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Document:\n\n{text[:MAX_TEXT_CHARS]}"},
        ],
        response_format={"type": "json_object"},
        temperature=0,
        max_tokens=1500,
    )
    raw = response.choices[0].message.content or "{}"
    return json.loads(raw)


# ─── Endpoints ─────────────────────────────────────────────────────────────

class ScrapeUrlRequest(BaseModel):
    url: str = Field(..., max_length=2048)


@router.post("/scrape-url", response_model=ParsedCompanyProfile)
async def scrape_url(req: ScrapeUrlRequest) -> ParsedCompanyProfile:
    """Fetch a company website and extract structured profile fields.

    Uses httpx to fetch the page HTML, extracts visible text via
    BeautifulSoup, then runs the same GPT-4o-mini extraction pipeline
    as the document parser.
    """
    import re

    import httpx
    from bs4 import BeautifulSoup

    url = req.url.strip()
    if not url.startswith(("http://", "https://")):
        url = f"https://{url}"

    # ── SSRF guard: reject loopback, private, and link-local addresses ──
    import ipaddress
    import socket
    from urllib.parse import urlparse

    parsed = urlparse(url)
    hostname = parsed.hostname
    if not hostname:
        raise HTTPException(status_code=422, detail="Invalid URL — no hostname found.")

    try:
        # Resolve hostname to IP(s) before connecting
        addr_infos = socket.getaddrinfo(hostname, None, socket.AF_UNSPEC, socket.SOCK_STREAM)
        for family, _, _, _, sockaddr in addr_infos:
            ip = ipaddress.ip_address(sockaddr[0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                raise HTTPException(
                    status_code=422,
                    detail="URL points to a private or internal address. Please provide a public URL.",
                )
    except socket.gaierror:
        raise HTTPException(
            status_code=422,
            detail="Could not resolve hostname. Please check the URL.",
        )

    # ── SSRF helper for redirect validation ──
    def _check_ip_safe(host: str) -> None:
        """Raise HTTPException if *host* resolves to a private/internal IP."""
        try:
            for _fam, _, _, _, _sa in socket.getaddrinfo(host, None, socket.AF_UNSPEC, socket.SOCK_STREAM):
                _ip = ipaddress.ip_address(_sa[0])
                if _ip.is_private or _ip.is_loopback or _ip.is_link_local or _ip.is_reserved:
                    raise HTTPException(
                        status_code=422,
                        detail="URL redirected to a private address. Please provide a direct public URL.",
                    )
        except socket.gaierror:
            pass  # Unresolvable redirect target — httpx will fail on its own

    # 1. Fetch the page (follow redirects, but validate final destination)
    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            max_redirects=5,
            timeout=15.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; GTMAdvisor/1.0)"},
        ) as client:
            resp = await client.get(url)
            resp.raise_for_status()

        # SSRF: validate the final URL after redirects
        final_host = urlparse(str(resp.url)).hostname
        if final_host and final_host != hostname:
            _check_ip_safe(final_host)

        # Guard: reject very large responses to prevent memory exhaustion
        MAX_RESPONSE_BYTES = 2 * 1024 * 1024  # 2 MB
        content_length = resp.headers.get("content-length")
        if content_length and int(content_length) > MAX_RESPONSE_BYTES:
            raise HTTPException(
                status_code=422,
                detail="Page is too large to process. Please upload a document instead.",
            )
        if len(resp.content) > MAX_RESPONSE_BYTES:
            raise HTTPException(
                status_code=422,
                detail="Page is too large to process. Please upload a document instead.",
            )
    except HTTPException:
        raise
    except httpx.HTTPStatusError as exc:
        logger.warning("url_scrape_http_error", url=url, status=exc.response.status_code)
        raise HTTPException(
            status_code=422,
            detail=f"Could not fetch URL (HTTP {exc.response.status_code}). Please check the URL and try again.",
        ) from exc
    except Exception as exc:
        logger.warning("url_scrape_network_error", url=url, error=str(exc))
        raise HTTPException(
            status_code=422,
            detail="Could not reach the provided URL. Please check it is publicly accessible or upload a document instead.",
        ) from exc

    # 2. Extract visible text from HTML
    html = resp.text
    soup = BeautifulSoup(html, "html.parser")

    # Remove script, style, nav, footer, header tags
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "svg", "iframe"]):
        tag.decompose()

    # Get text content
    text = soup.get_text(separator="\n", strip=True)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Also extract meta description and OG tags for richer context
    meta_parts: list[str] = []
    meta_desc = soup.find("meta", attrs={"name": "description"})
    if meta_desc and meta_desc.get("content"):
        meta_parts.append(f"Meta description: {meta_desc['content']}")
    og_desc = soup.find("meta", attrs={"property": "og:description"})
    if og_desc and og_desc.get("content"):
        meta_parts.append(f"OG description: {og_desc['content']}")
    og_title = soup.find("meta", attrs={"property": "og:title"})
    if og_title and og_title.get("content"):
        meta_parts.append(f"OG title: {og_title['content']}")
    title_tag = soup.find("title")
    if title_tag and title_tag.string:
        meta_parts.append(f"Page title: {title_tag.string.strip()}")

    enriched_text = ""
    if meta_parts:
        enriched_text = "\n".join(meta_parts) + "\n\n"
    enriched_text += text

    stripped = enriched_text.strip()
    if not stripped or len(stripped) < 50:
        raise HTTPException(
            status_code=422,
            detail="Could not extract meaningful text from this URL. The page may require JavaScript to render. Please upload a document instead.",
        )

    logger.info("url_scrape_requested", url=url, chars=len(stripped))

    # 3. LLM extraction (reuse same pipeline as document parser)
    try:
        extracted = await _llm_extract(stripped, system_prompt=WEB_SYSTEM_PROMPT)
    except Exception as exc:
        logger.warning("url_llm_extraction_failed", url=url, error=str(exc))
        raise HTTPException(
            status_code=502,
            detail="AI extraction failed. Please fill in the details manually.",
        ) from exc

    # 4. Sanitise industry
    industry = extracted.get("industry")
    if industry and industry not in VALID_INDUSTRIES:
        industry = "other"

    warning = None
    if len(stripped) > MAX_TEXT_CHARS:
        warning = f"Page is large ({len(stripped):,} characters). Only the first {MAX_TEXT_CHARS:,} characters were analysed."

    return ParsedCompanyProfile(
        company_name=extracted.get("company_name") or None,
        description=extracted.get("description") or None,
        industry=industry,
        value_proposition=extracted.get("value_proposition") or None,
        goals=_to_str_list(extracted.get("goals")),
        challenges=_to_str_list(extracted.get("challenges")),
        competitors=_to_str_list(extracted.get("competitors")),
        target_markets=_to_str_list(extracted.get("target_markets")),
        extracted_chars=len(stripped),
        extracted_text=stripped[:MAX_TEXT_CHARS],
        warning=warning,
    )


@router.post("/parse", response_model=ParsedCompanyProfile)
async def parse_document(file: UploadFile = File(...)) -> ParsedCompanyProfile:
    """
    Extract structured company profile fields from an uploaded PDF or text file.

    - Max 10 MB
    - Accepts .pdf and .txt
    - Returns null for fields that cannot be determined from the document
    """
    # 1. Read + size guard
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="File is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {MAX_FILE_BYTES // 1_048_576} MB.",
        )

    # 2. Extract text
    try:
        raw_text = _extract_text(
            data,
            filename=file.filename or "",
            content_type=file.content_type or "",
        )
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    stripped = raw_text.strip()
    if not stripped:
        raise HTTPException(
            status_code=422,
            detail=(
                "No readable text found in this document. "
                "If it's a scanned PDF, please type or paste the content instead."
            ),
        )

    logger.info(
        "document_parse_requested",
        filename=file.filename,
        chars=len(stripped),
    )

    # 3. LLM extraction
    try:
        extracted = await _llm_extract(stripped)
    except Exception as exc:
        logger.warning("document_llm_extraction_failed", error=str(exc))
        raise HTTPException(
            status_code=502,
            detail=(
                f"AI extraction failed ({exc}). "
                "Please fill in the details manually or try a different file."
            ),
        ) from exc

    # 4. Sanitise industry value
    industry = extracted.get("industry")
    if industry and industry not in VALID_INDUSTRIES:
        industry = "other"

    warning = None
    if len(stripped) > MAX_TEXT_CHARS:
        warning = (
            f"Document is long ({len(stripped):,} characters). "
            f"Only the first {MAX_TEXT_CHARS:,} characters were analysed."
        )

    try:
        return ParsedCompanyProfile(
            company_name=extracted.get("company_name") or None,
            description=extracted.get("description") or None,
            industry=industry,
            value_proposition=extracted.get("value_proposition") or None,
            goals=_to_str_list(extracted.get("goals")),
            challenges=_to_str_list(extracted.get("challenges")),
            competitors=_to_str_list(extracted.get("competitors")),
            target_markets=_to_str_list(extracted.get("target_markets")),
            extracted_chars=len(stripped),
            extracted_text=stripped[:MAX_TEXT_CHARS],
            warning=warning,
        )
    except Exception as exc:
        logger.warning("document_profile_construction_failed", error=str(exc))
        raise HTTPException(
            status_code=502,
            detail="AI returned an unexpected response format. Please fill in the details manually.",
        ) from exc
